import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


async def generate_visit_docx(visit_data: dict, doctor_settings: dict = None) -> bytes:
    """Генерация Word для одного приёма"""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Шапка
    header_text = ""
    if doctor_settings:
        header_text = doctor_settings.get("doc_header", "") or ""
        doctor_name = doctor_settings.get("doc_doctor_name", "") or ""
        doctor_contacts = doctor_settings.get("doc_doctor_contacts", "") or ""
    else:
        doctor_name = ""
        doctor_contacts = ""

    if header_text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(74, 144, 217)

    if doctor_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(doctor_name)

    if doctor_contacts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(doctor_contacts)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph("_" * 60)

    # Данные пациента
    doc.add_heading("Данные пациента", level=2)
    doc.add_paragraph(f"Владелец: {visit_data.get('owner_name', '')}")
    doc.add_paragraph(
        f"Питомец: {visit_data.get('pet_name', '')} "
        f"({visit_data.get('pet_species', '')}"
        f"{', ' + visit_data.get('pet_breed', '') if visit_data.get('pet_breed') else ''})"
    )

    visit_date = visit_data.get("visit_date", "")
    if isinstance(visit_date, datetime):
        visit_date = visit_date.strftime("%d.%m.%Y %H:%M")
    doc.add_paragraph(f"Дата: {visit_date}")

    visit_type = "Первичный" if visit_data.get("visit_type") == "primary" else "Повторный"
    doc.add_paragraph(f"Тип: {visit_type}")

    if visit_data.get("weight"):
        doc.add_paragraph(f"Вес: {visit_data['weight']} кг")

    # Анамнез
    if visit_data.get("anamnesis"):
        doc.add_heading("Анамнез", level=2)
        doc.add_paragraph(visit_data["anamnesis"])

    # Рекомендации
    if visit_data.get("recommendations"):
        doc.add_heading("Рекомендации", level=2)
        doc.add_paragraph(visit_data["recommendations"])

    # Примечания
    if visit_data.get("notes"):
        doc.add_heading("Примечания", level=2)
        doc.add_paragraph(visit_data["notes"])

    # Подпись
    doc.add_paragraph("_" * 60)
    if doctor_settings:
        footer = doctor_settings.get("doc_footer", "") or ""
        if footer:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(footer)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


async def generate_epicrisis_docx(visits_data: list, pet_data: dict, doctor_settings: dict = None) -> bytes:
    """Генерация эпикриза Word — все приёмы в 1 файл"""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Шапка
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ЭПИКРИЗ")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(74, 144, 217)

    if doctor_settings:
        header = doctor_settings.get("doc_header", "") or ""
        if header:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(header)

    doc.add_paragraph("_" * 60)

    # Данные пациента
    doc.add_heading("Пациент", level=2)
    doc.add_paragraph(f"Владелец: {pet_data.get('owner_name', '')}")
    doc.add_paragraph(
        f"Питомец: {pet_data.get('name', '')} "
        f"({pet_data.get('species', '')}"
        f"{', ' + pet_data.get('breed', '') if pet_data.get('breed') else ''})"
    )
    doc.add_paragraph(f"Возраст: {pet_data.get('age', '')}")

    doc.add_paragraph("_" * 60)

    # Приёмы
    doc.add_heading(f"История приёмов ({len(visits_data)})", level=2)

    for i, visit in enumerate(visits_data, 1):
        visit_date = visit.get("visit_date", "")
        if isinstance(visit_date, datetime):
            visit_date = visit_date.strftime("%d.%m.%Y %H:%M")

        visit_type = "Первичный" if visit.get("visit_type") == "primary" else "Повторный"
        doc.add_heading(f"Приём #{i} — {visit_date} ({visit_type})", level=3)

        if visit.get("weight"):
            doc.add_paragraph(f"Вес: {visit['weight']} кг")

        if visit.get("anamnesis"):
            p = doc.add_paragraph()
            p.add_run("Анамнез: ").bold = True
            p.add_run(visit["anamnesis"])

        if visit.get("recommendations"):
            p = doc.add_paragraph()
            p.add_run("Рекомендации: ").bold = True
            p.add_run(visit["recommendations"])

        if visit.get("notes"):
            p = doc.add_paragraph()
            p.add_run("Примечания: ").bold = True
            p.add_run(visit["notes"])

        doc.add_paragraph("—" * 40)

    # Подпись
    if doctor_settings:
        footer = doctor_settings.get("doc_footer", "") or ""
        if footer:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(footer)
            run.font.size = Pt(9)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
